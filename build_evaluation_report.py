#!/usr/bin/env python3
"""Render the privacy-safe Markdown evaluation report as an upload-ready DOCX.

The document is generated from ``EVALUATION.md`` so that the downloadable
report and the repository version remain in sync. It contains no source
prospectus text or annotation values.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "EVALUATION.md"
OUTPUT = ROOT / "output" / "PII_Redaction_Evaluation_Report.docx"
BLUE = RGBColor(38, 89, 151)
SLATE = RGBColor(66, 75, 88)


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_text(cell, value: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.size = Pt(8.5)


def _add_inline(paragraph, value: str) -> None:
    """Render the small Markdown subset used by EVALUATION.md."""
    tokens = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", value)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(8.7)
        else:
            paragraph.add_run(token)


def _table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|[\s|:-]+\|", line.strip()))


def _add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Light Shading Accent 1"
    table.autofit = True
    for index, value in enumerate(rows[0]):
        _set_cell_text(table.rows[0].cells[index], value, bold=True)
        _shade(table.rows[0].cells[index], "D9EAF7")
    for values in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            _set_cell_text(cells[index], value.replace("**", ""))
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["No Spacing"]
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(7.8)
        run.font.color.rgb = SLATE
        if index < len(lines) - 1:
            run.add_break()


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = SLATE
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for level, size in ((1, 16), (2, 12.5)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(12 if level == 1 else 9)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("ScalePII  |  Evaluation Strategy & Metrics")
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(104, 116, 132)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Privacy-safe report — no raw prospectus or annotation values included")
    footer_run.font.size = Pt(7.5)
    footer_run.font.color.rgb = RGBColor(104, 116, 132)


def render(source: Path, output: Path) -> None:
    document = Document()
    _configure_document(document)
    document.core_properties.title = "ScalePII — Evaluation Strategy & Metrics"
    document.core_properties.author = "Ankan Ghosh"
    document.core_properties.subject = "PII redaction evaluation report"

    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    title_added = False
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run("ScalePII — " + line[2:])
            run.font.name = "Aptos Display"
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = BLUE
            title_added = True
            index += 1
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            index += 1
            continue
        if line.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            _add_code_block(document, code_lines)
            index += 1
            continue
        if line.startswith("|"):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].startswith("|"):
                if not _is_separator(lines[index]):
                    rows.append(_table_row(lines[index]))
                index += 1
            _add_table(document, rows)
            continue

        paragraph_lines = [line.strip()]
        index += 1
        while (
            index < len(lines)
            and lines[index].strip()
            and not lines[index].startswith(("#", "|", "```"))
        ):
            paragraph_lines.append(lines[index].strip())
            index += 1
        paragraph = document.add_paragraph()
        _add_inline(paragraph, " ".join(paragraph_lines))

    if not title_added:
        raise RuntimeError(f"missing title in {source}")
    output.parent.mkdir(exist_ok=True)
    document.save(output)
    print(f"Wrote {output.relative_to(ROOT)} ({output.stat().st_size:,} bytes)")


if __name__ == "__main__":
    render(SOURCE, OUTPUT)
