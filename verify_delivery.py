#!/usr/bin/env python3
"""Independently verify the final redacted DOCX and its audit artefacts.

This verifier intentionally checks the *serialized package*, not merely the
text displayed by Word. That matters because a DOCX can retain source PII in
hidden hyperlink instructions or embedded media even after visible text has
been changed.

It proves only the stated, mechanical properties: the final package is valid;
its structure is preserved; all source values detected by the current pipeline
are absent from its XML; redacted image parts are neutral placeholders; and the
generated audit artefacts do not carry cleartext source values. It cannot prove
that an unlabelled value was never missed by the detector.
"""

from __future__ import annotations

import argparse
import csv
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree

from docx.oxml.ns import qn

from piiredact.allowlist import load_glossary
from piiredact.detectors import build_detectors
from piiredact.docx_io import load_blocks
from piiredact.pipeline import harvest, harvest_glossary
from piiredact.spans import resolve_overlaps
from redact import load_nlp


XML_SUFFIXES = (".xml", ".rels")
FIELD_TAGS = {qn("w:instrText"), qn("w:delInstrText")}


def _package_xml(path: Path) -> tuple[dict[str, bytes], bytes]:
    with zipfile.ZipFile(path) as package:
        if package.testzip() is not None:
            raise AssertionError(f"invalid DOCX package: {path}")
        parts = {
            name: package.read(name)
            for name in package.namelist()
            if name.endswith(XML_SUFFIXES)
        }
    return parts, b"".join(parts.values())


def _bold_run_count(path: Path) -> int:
    total = 0
    with zipfile.ZipFile(path) as package:
        for name in package.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = ElementTree.fromstring(package.read(name))
            for run in root.iter(qn("w:r")):
                properties = run.find(qn("w:rPr"))
                if properties is None:
                    continue
                bold = properties.find(qn("w:b"))
                if bold is not None and bold.get(qn("w:val"), "true").lower() not in {
                    "0",
                    "false",
                    "off",
                }:
                    total += 1
    return total


def _field_instruction_values(xml_parts: dict[str, bytes], detectors) -> list[str]:
    """Return detected source values inside hidden Word field instructions."""
    values: list[str] = []
    for data in xml_parts.values():
        root = ElementTree.fromstring(data)
        for node in root.iter():
            if node.tag not in FIELD_TAGS or not node.text:
                continue
            candidates = []
            for detector in detectors:
                candidates.extend(detector.find(node.text))
            values.extend(span.text for span in resolve_overlaps(candidates))
    return values


def _read_jsonl(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line]


def _assert_no_source_values(output_xml: bytes, values: set[str], label: str) -> None:
    leaked = [value for value in values if value.encode("utf-8") in output_xml]
    if leaked:
        # Never print source values in a privacy verifier. A hash permits
        # correlation with the privacy-safe audit files.
        raise AssertionError(f"{label}: {len(leaked)} source values remain in output XML")


def verify(input_path: Path, output_path: Path, model: str) -> dict[str, object]:
    source_document, source_blocks = load_blocks(str(input_path))
    redacted_document, redacted_blocks = load_blocks(str(output_path))
    source_parts, source_xml = _package_xml(input_path)
    _output_parts, output_xml = _package_xml(output_path)

    nlp = load_nlp(model)
    if nlp is None:
        raise RuntimeError("spaCy model is required for verification of the submitted run")

    load_glossary(harvest_glossary(source_document))
    gazetteers, _ = harvest(source_blocks, nlp, document=source_document)
    detectors = build_detectors(
        nlp=None,
        person_terms=gazetteers.persons,
        org_terms=gazetteers.orgs,
        known_dins=gazetteers.dins,
    )

    visible_spans = []
    for block in source_blocks:
        candidates = []
        for detector in detectors:
            candidates.extend(detector.find(block.text))
        visible_spans.extend(resolve_overlaps(candidates))
    visible_values = {span.text for span in visible_spans}
    field_values = _field_instruction_values(source_parts, detectors)
    _assert_no_source_values(output_xml, visible_values, "visible-text leak check")
    _assert_no_source_values(output_xml, set(field_values), "hidden-field leak check")

    source_emails = set(
        re.findall(rb"(?i)[a-z0-9._%+\-]+@[a-z0-9.\-]+\.[a-z]{2,}", source_xml)
    )
    if any(email in output_xml for email in source_emails):
        raise AssertionError("source email values remain in output XML")

    expected_structure = {
        "paragraphs": len(source_document.paragraphs),
        "tables": len(source_document.tables),
        "sections": len(source_document.sections),
        "text_blocks": len(source_blocks),
        "bold_runs": _bold_run_count(input_path),
    }
    actual_structure = {
        "paragraphs": len(redacted_document.paragraphs),
        "tables": len(redacted_document.tables),
        "sections": len(redacted_document.sections),
        "text_blocks": len(redacted_blocks),
        "bold_runs": _bold_run_count(output_path),
    }
    if actual_structure != expected_structure:
        raise AssertionError("document structure changed during redaction")

    asset_dir = Path(__file__).parent / "assets"
    placeholders = {
        ".png": (asset_dir / "redacted-image-placeholder.png").read_bytes(),
        ".jpeg": (asset_dir / "redacted-image-placeholder.jpeg").read_bytes(),
        ".jpg": (asset_dir / "redacted-image-placeholder.jpeg").read_bytes(),
    }
    with zipfile.ZipFile(output_path) as package:
        media_names = sorted(name for name in package.namelist() if name.startswith("word/media/"))
        mismatches = [
            name
            for name in media_names
            if Path(name).suffix.lower() not in placeholders
            or package.read(name) != placeholders[Path(name).suffix.lower()]
        ]
    if mismatches:
        raise AssertionError(f"{len(mismatches)} embedded images are not neutral placeholders")

    output_dir = output_path.parent
    mapping_rows = list(
        csv.DictReader((output_dir / "mapping.csv").open(encoding="utf-8", newline=""))
    )
    detection_rows = _read_jsonl(output_dir / "detections.jsonl")
    field_rows = json.loads((output_dir / "field_code_redactions.json").read_text(encoding="utf-8"))
    media_rows = json.loads((output_dir / "media_redactions.json").read_text(encoding="utf-8"))
    run_summary = json.loads(
        (Path("reports") / "run_summary.json").read_text(encoding="utf-8")
    )

    expected_mapping_fields = {"entity_type", "source_sha256", "replacement"}
    if not mapping_rows or any(set(row) != expected_mapping_fields for row in mapping_rows):
        raise AssertionError("mapping.csv has an unexpected privacy-audit schema")
    if any("text" in row or "source_sha256" not in row for row in detection_rows):
        raise AssertionError("detections.jsonl carries cleartext source data or lacks hashes")
    if any("text" in row or "source_sha256" not in row for row in field_rows):
        raise AssertionError("field-code audit carries cleartext source data or lacks hashes")
    if len(detection_rows) != run_summary["replacements"]:
        raise AssertionError("visible detection audit count does not match run summary")
    if len(field_rows) != run_summary["field_code_values_redacted"]:
        raise AssertionError("field-code audit count does not match run summary")
    if len(media_rows) != run_summary["image_assets_redacted"] or len(media_rows) != len(media_names):
        raise AssertionError("media audit count does not match package or run summary")
    if len(mapping_rows) != run_summary["stable_mappings"]:
        raise AssertionError("mapping count does not match run summary")

    report_paths = sorted(Path("reports").glob("evaluation*.json"))
    if any('"text"' in path.read_text(encoding="utf-8") for path in report_paths):
        raise AssertionError("evaluation report JSON carries cleartext source values")

    return {
        "package_valid": True,
        "structure": expected_structure,
        "detected_visible_occurrences_checked": len(visible_spans),
        "detected_visible_unique_values_checked": len(visible_values),
        "detected_field_occurrences_checked": len(field_values),
        "source_email_values_checked": len(source_emails),
        "neutralized_image_assets": len(media_names),
        "audit_rows": {
            "mapping": len(mapping_rows),
            "visible_text": len(detection_rows),
            "hidden_fields": len(field_rows),
            "media": len(media_rows),
        },
        "detected_by_type": dict(sorted(Counter(span.entity_type for span in visible_spans).items())),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", default="input.docx")
    parser.add_argument("--output", default="output/Red Herring Prospectus - REDACTED.docx")
    parser.add_argument("--model", default="en_core_web_sm")
    args = parser.parse_args()
    print(json.dumps(verify(Path(args.input), Path(args.output), args.model), indent=2))


if __name__ == "__main__":
    main()
