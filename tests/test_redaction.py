#!/usr/bin/env python3
"""Regression tests for silent or high-risk redaction failure modes.

The suite covers Word-run fragmentation, Word field instructions, image
payloads, false-positive controls and every PII type required by the assignment.
Several tests were added after defects observed during development; others are
synthetic integration checks for required functionality.

    python -m unittest discover tests -v
"""

from __future__ import annotations

import sys
import unittest
from base64 import b64decode
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.parts.image import ImagePart  # noqa: E402

from piiredact.allowlist import (  # noqa: E402
    has_legal_entity_marker,
    is_only_generic,
    is_person_noise,
    is_redactable_org,
    load_glossary,
)
from piiredact.detectors import (  # noqa: E402
    AddressDetector,
    CorporateEntityDetector,
    CreditCardDetector,
    DateOfBirthDetector,
    EmailDetector,
    GazetteerDetector,
    PhoneDetector,
    WebsiteDetector,
    _luhn_ok,
)
from piiredact.docx_io import (  # noqa: E402
    apply_to_block,
    load_blocks,
    redact_field_instructions,
)
from piiredact.fakes import FakeFactory  # noqa: E402
from piiredact.media import redact_embedded_images  # noqa: E402
from piiredact.pipeline import redact  # noqa: E402
from piiredact.spans import Span, resolve_overlaps  # noqa: E402


def texts(detector, value):
    return [span.text for span in detector.find(value)]


class TestRunFragmentation(unittest.TestCase):
    """The core .docx problem: entities split across formatting runs."""

    def _document_with_runs(self, pieces):
        path = Path("/tmp/_piiredact_test.docx")
        document = Document()
        paragraph = document.add_paragraph()
        for piece in pieces:
            run = paragraph.add_run(piece)
            run.bold = True
        document.save(path)
        return path

    def test_entity_split_across_runs_is_replaced(self):
        # Synthetic email deliberately split across Word formatting runs.
        path = self._document_with_runs(
            ["E-mail: ", "", "contact@", "acme", ".example.org", "; Website"]
        )
        _, blocks = load_blocks(str(path))
        block = blocks[0]
        self.assertIn("contact@acme.example.org", block.text)

        spans = EmailDetector().find(block.text)
        self.assertEqual(len(spans), 1, "email spanning runs must be found once")

        apply_to_block(block, [(spans[0], "fake@example.com")])
        _, reloaded = load_blocks(str(path.parent / path.name))
        # Re-read from the in-memory runs rather than disk state.
        rebuilt = "".join(run.text for _s, _e, run in block.layout)
        self.assertNotIn("acme.example", rebuilt)
        self.assertIn("fake@example.com", rebuilt)

    def test_formatting_survives_replacement(self):
        path = self._document_with_runs(["Contact ", "person@example.org", " now"])
        _, blocks = load_blocks(str(path))
        block = blocks[0]
        spans = EmailDetector().find(block.text)
        apply_to_block(block, [(spans[0], "x@example.com")])
        for _s, _e, run in block.layout:
            self.assertTrue(run._run.bold, "run properties must not be disturbed")

    def test_linked_headers_are_processed_once(self):
        path = Path("/tmp/_piiredact_linked_headers.docx")
        document = Document()
        document.add_paragraph("Body text")
        document.sections[0].header.paragraphs[0].text = "Contact fake@example.com"
        document.add_section()
        document.save(path)

        _, blocks = load_blocks(str(path))
        header_blocks = [block for block in blocks if block.location.startswith("header:")]
        self.assertEqual(len(header_blocks), 1)
        self.assertEqual(header_blocks[0].text, "Contact fake@example.com")


class TestEmbeddedImageRedaction(unittest.TestCase):
    """Embedded assets can contain PII even when the text layer is clean."""

    def test_all_embedded_images_are_replaced_and_document_stays_valid(self):
        image_path = Path("/tmp/_piiredact_test_source_image.png")
        # A tiny synthetic PNG keeps the test independent of the supplied
        # prospectus, so the test suite can run from the submission package.
        image_path.write_bytes(
            b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4z8DwHwAFgAI/ScL9xQAAAABJRU5ErkJggg=="
            )
        )
        source_path = Path("/tmp/_piiredact_images_source.docx")
        document = Document()
        document.add_picture(str(image_path))
        document.save(source_path)

        image_parts = [
            part for part in document.part.package.parts if isinstance(part, ImagePart)
        ]
        original_blobs = {part.partname: part.blob for part in image_parts}

        audit_rows = redact_embedded_images(document)

        self.assertEqual(len(audit_rows), len(image_parts))
        self.assertTrue(
            all(part.blob != original_blobs[part.partname] for part in image_parts)
        )
        self.assertTrue(
            all(row["action"] == "replaced_with_neutral_placeholder" for row in audit_rows)
        )

        path = Path("/tmp/_piiredact_images_redacted.docx")
        document.save(path)
        reloaded = Document(path)
        reloaded_parts = [
            part for part in reloaded.part.package.parts if isinstance(part, ImagePart)
        ]
        self.assertEqual(len(reloaded_parts), len(image_parts))


class TestFieldInstructionRedaction(unittest.TestCase):
    """A visible fake must not leave a real mailto/URL target in DOCX XML."""

    def test_hyperlink_field_targets_are_redacted(self):
        path = Path("/tmp/_piiredact_field_codes.docx")
        document = Document()
        paragraph = document.add_paragraph()
        for instruction in (
            ' HYPERLINK "mailto:real.person@example.org"',
            ' HYPERLINK "https://www.real-company.example.org/"',
        ):
            field = OxmlElement("w:instrText")
            field.set(qn("xml:space"), "preserve")
            field.text = instruction
            paragraph._p.append(field)

        audit_rows = redact_field_instructions(
            document,
            [EmailDetector(), WebsiteDetector()],
            FakeFactory(seed=99),
        )
        self.assertEqual(len(audit_rows), 2)
        self.assertTrue(all("source_sha256" in row for row in audit_rows))
        self.assertTrue(all("text" not in row for row in audit_rows))

        document.save(path)
        package_bytes = path.read_bytes()
        self.assertNotIn(b"real.person@example.org", package_bytes)
        self.assertNotIn(b"www.real-company.example.org", package_bytes)


class TestOverlapResolution(unittest.TestCase):
    def test_priority_beats_length(self):
        short_high = Span(0, 5, "aaaaa", "ORG", "regex", priority=80)
        long_low = Span(0, 40, "a" * 40, "ADDRESS", "markers", priority=40)
        kept = resolve_overlaps([long_low, short_high])
        self.assertEqual([s.entity_type for s in kept], ["ORG"])

    def test_non_overlapping_spans_all_survive(self):
        spans = [
            Span(0, 5, "aaaaa", "ORG", "regex", priority=80),
            Span(10, 20, "b" * 10, "ADDRESS", "markers", priority=40),
        ]
        self.assertEqual(len(resolve_overlaps(spans)), 2)


class TestClosedFormDetectors(unittest.TestCase):
    def test_luhn_rejects_ordinary_financial_figures(self):
        # A prospectus is full of long numbers; without Luhn every one of
        # them is a "credit card".
        self.assertFalse(_luhn_ok("1234567890123456"))
        self.assertTrue(_luhn_ok("4539578763621486"))

    def test_credit_card_detector_ignores_share_counts(self):
        detector = CreditCardDetector()
        self.assertEqual(texts(detector, "56,818,200 Equity Shares of face value"), [])
        self.assertEqual(texts(detector, "Card 4539 5787 6362 1486 on file"),
                         ["4539 5787 6362 1486"])

    def test_phone_requires_prefix_or_label(self):
        detector = PhoneDetector()
        self.assertEqual(texts(detector, "aggregating up to 4200000000 million"), [])
        self.assertEqual(texts(detector, "Telephone: + 91 98765 43210"), ["+ 91 98765 43210"])

    def test_date_of_birth_needs_a_birth_cue(self):
        detector = DateOfBirthDetector()
        self.assertEqual(texts(detector, "a resolution dated December 11, 2024"), [])
        self.assertEqual(texts(detector, "Date of Birth: March 3, 1961"), ["March 3, 1961"])

    def test_ssn_and_ip_address_are_detected(self):
        from piiredact.detectors import IPAddressDetector, SSNDetector

        self.assertEqual(texts(SSNDetector(), "SSN: 123-45-6789"), ["123-45-6789"])
        self.assertEqual(texts(IPAddressDetector(), "Host: 203.0.113.7"), ["203.0.113.7"])


class TestMinimumTypeIntegration(unittest.TestCase):
    """Exercise every PII category required by the assignment end to end."""

    def test_all_required_types_are_replaced_and_do_not_survive_in_package_xml(self):
        source_path = Path("/tmp/_piiredact_minimum_types_source.docx")
        output_path = Path("/tmp/_piiredact_minimum_types_output.docx")
        document = Document()
        paragraph = document.add_paragraph()
        # Deliberately split the values across runs: real Word documents do
        # this, and the production pipeline must flatten before detecting.
        values = [
            "Contact Person: Anita Rao; ",
            "Acme Private Limited; ",
            "anita.rao@example.org; ",
            "+91 98765 43210; ",
            "12 Example Road, Example City – 411001, Example State, India; ",
            "SSN: 123-45-6789; ",
            "Card: 4539 5787 6362 1486; ",
            "Date of Birth: March 3, 1961; ",
            "IP: 203.0.113.7; ",
            "www.acme.example.org",
        ]
        for value in values:
            paragraph.add_run(value)
        document.save(source_path)

        result = redact(str(source_path), str(output_path), nlp=None, seed=88)
        expected_types = {
            "PERSON", "ORG", "EMAIL", "PHONE", "ADDRESS", "SSN",
            "CREDIT_CARD", "DATE_OF_BIRTH", "IP_ADDRESS", "WEBSITE",
        }
        self.assertTrue(expected_types.issubset(result["by_type"]))

        import zipfile

        with zipfile.ZipFile(output_path) as package:
            output_xml = b"".join(
                package.read(name)
                for name in package.namelist()
                if name.endswith((".xml", ".rels"))
            )
        for value in values:
            candidate = value.strip(" ;")
            if candidate.startswith(("Contact Person:", "SSN:", "Card:", "Date of Birth:", "IP:")):
                candidate = candidate.split(": ", 1)[1]
            self.assertNotIn(candidate.encode("utf-8"), output_xml)


class TestAddressDetector(unittest.TestCase):
    def setUp(self):
        self.detector = AddressDetector()

    def test_registration_number_is_not_a_pin_code(self):
        # A registration number ending in six digits must not anchor an address span
        # across the whole sentence, suppressing the PERSON inside it.
        found = texts(
            self.detector,
            "the independent chartered engineer appointed by our Company, "
            "namely, Rohan Mehta bearing registration number M-123456",
        )
        self.assertEqual(found, [])

    def test_indian_abbreviations_do_not_truncate_the_span(self):
        # "S. no." was being treated as a sentence boundary, so the house
        # number and building name survived while the rest was replaced.
        found = texts(
            self.detector,
            "S. no. 245/ 104, Sunrise Residency, Orchard Society, lane no. 3 "
            "Maple Road, Central District, Example City – 411004 Example State, India",
        )
        self.assertEqual(len(found), 1)
        self.assertTrue(found[0].startswith("S. no. 245/ 104, Sunrise Residency"))

    def test_address_excludes_the_company_that_sits_at_it(self):
        found = texts(
            self.detector,
            "Apex Securities Limited Apex Venture House Market Road "
            "Central District, Example City – 400025 Example State, India",
        )
        self.assertEqual(len(found), 1)
        self.assertNotIn("Apex Securities Limited", found[0])

    def test_address_line_without_a_pin_is_still_found(self):
        found = texts(
            self.detector,
            "201, Tower-2, Horizon Business Centre Off Lakeside Avenue, West End",
        )
        self.assertEqual(len(found), 1)


class TestOrganisationPolicy(unittest.TestCase):
    def setUp(self):
        load_glossary({"Bid Amount", "Anchor Investor", "Key Managerial Personnel"})

    def test_regulators_and_statutes_are_kept(self):
        for value in ("SEBI", "Companies Act, 2013", "BSE", "Reserve Bank of India"):
            self.assertFalse(is_redactable_org(value, set()), value)

    def test_named_legal_entities_are_redacted(self):
        for value in ("Acme Metals Limited", "Harbor & Co. LLP",
                      "Northern Alloy AB", "Summit Family Trust"):
            self.assertTrue(is_redactable_org(value, set()), value)

    def test_glossary_terms_are_not_entities(self):
        self.assertFalse(is_redactable_org("Bid Amount", set()))

    def test_weak_markers_only_count_in_trailing_position(self):
        # "…capital goods" must not read as a company because of "goods";
        # "Harbor Associates" must, because "Associates" ends the name.
        self.assertFalse(has_legal_entity_marker("export promotion capital goods"))
        self.assertTrue(has_legal_entity_marker("Harbor Associates"))

    def test_bare_corporate_boilerplate_is_not_an_entity(self):
        for value in ("Private Limited", "Family Trust", "Company"):
            self.assertTrue(is_only_generic(value), value)

    def test_corporate_detector_finds_what_ner_misses(self):
        found = texts(
            CorporateEntityDetector(), "(Formerly Harbor Registry Private Limited)"
        )
        self.assertIn("Harbor Registry Private Limited", found)


class TestPersonPolicy(unittest.TestCase):
    def setUp(self):
        load_glossary({"Bid Amount", "Floor Price"})

    def test_domain_vocabulary_is_not_a_person(self):
        for value in ("Bid Amount", "Floor Price", "Key Managerial Personnel",
                      "Mutual Funds", "Gram Jyoti", "ISO 9001:2015"):
            self.assertTrue(is_person_noise(value, strict=True), value)

    def test_real_names_are_accepted(self):
        for value in ("Arjun Mehta", "Maya Patel", "Anita Rao"):
            self.assertFalse(is_person_noise(value, strict=True), value)

    def test_entity_suffix_disqualifies_a_person(self):
        self.assertTrue(
            is_person_noise("Northpoint Industrial Park VI Private Limited", strict=True)
        )


class TestGazetteer(unittest.TestCase):
    def test_punctuation_variants_still_match(self):
        # Learned as "Alder & Finch LLP"; some documents also write
        # "Alder & Finch, LLP". Literal matching would miss the second.
        detector = GazetteerDetector("ORG", {"Alder & Finch LLP"})
        self.assertEqual(
            texts(detector, "consent from Alder & Finch, LLP, Chartered Accountants"),
            ["Alder & Finch, LLP"],
        )

    def test_matching_is_case_insensitive(self):
        detector = GazetteerDetector("PERSON", {"Arjun Mehta"})
        self.assertEqual(len(detector.find("OUR PROMOTERS: ARJUN MEHTA")), 1)

    def test_empty_gazetteer_is_safe(self):
        self.assertEqual(GazetteerDetector("ORG", set()).find("anything"), [])


class TestFakeFactory(unittest.TestCase):
    def setUp(self):
        self.factory = FakeFactory(seed=20251210)

    def test_mapping_is_stable_within_a_run(self):
        first = self.factory.fake_for("PERSON", "Arjun Mehta")
        second = self.factory.fake_for("PERSON", "Arjun Mehta")
        self.assertEqual(first, second)

    def test_mapping_is_reproducible_across_runs(self):
        other = FakeFactory(seed=20251210)
        self.assertEqual(
            self.factory.fake_for("PERSON", "Karan Singh"),
            other.fake_for("PERSON", "Karan Singh"),
        )

    def test_casing_is_preserved(self):
        fake = self.factory.fake_for("PERSON", "ARJUN MEHTA")
        self.assertEqual(fake, fake.upper())

    def test_whitespace_variants_share_one_fake(self):
        # "+ 91 20 4505 3237" and "+91 20 45053237" are the same number and
        # must not become two different fake numbers.
        a = self.factory.fake_for("PHONE", "+ 91 20 4505 3237")
        b = self.factory.fake_for("PHONE", "+ 91 20 4505  3237")
        self.assertEqual(a, b)

    def test_phone_punctuation_variants_share_one_fake(self):
        a = self.factory.fake_for("PHONE", "+ 91 (20) 6729 5100")
        b = self.factory.fake_for("PHONE", "+91 20 6729 5100")
        self.assertEqual(a, b)

    def test_organisation_punctuation_variants_share_one_fake(self):
        a = self.factory.fake_for("ORG", "Alder & Finch LLP")
        b = self.factory.fake_for("ORG", "Alder & Finch, LLP")
        self.assertEqual(a, b)

    def test_organisation_fake_preserves_legal_form(self):
        fake = self.factory.fake_for("ORG", "Northstar Materials LLC")
        self.assertTrue(fake.endswith(" LLC"))

    def test_mapping_hashes_the_source_value(self):
        source = "Alder & Finch LLP"
        self.factory.fake_for("ORG", source)
        row = self.factory.mapping_rows()[0]
        self.assertNotIn("original", row)
        self.assertNotIn(source, row.values())
        self.assertEqual(len(row["source_sha256"]), 64)

    def test_website_fake_preserves_url_scheme(self):
        factory = FakeFactory(seed=11)
        self.assertTrue(
            factory.fake_for("WEBSITE", "https://company.example.org").startswith("https://")
        )
        self.assertTrue(
            factory.fake_for("WEBSITE", "http://company.example.org").startswith("http://")
        )

    def test_fake_does_not_contain_the_original(self):
        fake = self.factory.fake_for("EMAIL", "contact@acme.example.org")
        self.assertNotIn("acme", fake)
        self.assertTrue(fake.endswith("@example.com"))

    def test_distinct_values_get_distinct_fakes(self):
        names = ["Aarav Sharma", "Maya Kapoor", "Vikram Joshi"]
        fakes = {self.factory.fake_for("PERSON", n) for n in names}
        self.assertEqual(len(fakes), len(names))


if __name__ == "__main__":
    unittest.main(verbosity=2)
